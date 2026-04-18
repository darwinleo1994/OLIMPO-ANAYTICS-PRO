import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
import numpy as np
import os
import statsmodels.api as sm
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# 1. Configuración de Marca y Estética Pro
st.set_page_config(page_title="Olimpo Analytics Pro", page_icon="🏛️", layout="wide")

st.markdown("""
    <style>
    #MainMenu {visibility: hidden;} footer {visibility: hidden;} header {visibility: hidden;}
    body { -webkit-user-select: none; user-select: none; }
    .stApp { background-color: #F8F9F9; color: #212F3D; }
    .main-title { color: #1B4F72; font-size: 45px; font-weight: bold; text-align: center; margin-bottom: 0px; }
    .section-header { color: #2874A6; border-bottom: 2px solid #AED6F1; padding-bottom: 5px; margin-top: 20px;}
    .interpret-box { background-color: #E8F6F3; padding: 20px; border-radius: 10px; border-left: 8px solid #1ABC9C; margin-top: 10px; }
    .edu-note { background-color: #FEF9E7; padding: 15px; border-radius: 5px; font-size: 15px; color: #7D6608; margin-bottom: 15px; border-left: 4px solid #F4D03F;}
    .premium-lock { background-color: #FDEDEC; padding: 15px; border-radius: 8px; border: 2px solid #E74C3C; color: #943126; font-weight: bold; }
    .step-box { background-color: #FFFFFF; padding: 20px; border-radius: 8px; box-shadow: 0px 2px 5px rgba(0,0,0,0.05); margin-bottom: 20px;}
    .plan-box { border: 1px solid #D5D8DC; padding: 10px; border-radius: 5px; margin-bottom: 10px; background-color: #FDFEFE; }
    .plan-vip { border: 2px solid #F39C12; padding: 10px; border-radius: 5px; margin-bottom: 10px; background-color: #FEF9E7; }
    </style>
""", unsafe_allow_html=True)

def limpiar_pantalla():
    for key in list(st.session_state.keys()):
        del st.session_state[key]

# --- BARRA LATERAL (PLANES Y ACCESO) ---
if os.path.exists("logo.png"):
    st.sidebar.image("logo.png", width=160)
else:
    st.sidebar.markdown("<h1 style='text-align: center;'>🏛️</h1>", unsafe_allow_html=True)

st.sidebar.markdown("### 🔑 Acceso Corporativo")
key_ingresada = st.sidebar.text_input("Licencia de Usuario:", type="password")

SOCIOS_OLIMPO = ["37219873", "37219874", "37219875", "37219876", "37219877"]
ES_PREMIUM = key_ingresada in SOCIOS_OLIMPO

if not ES_PREMIUM:
    st.sidebar.warning("⚠️ MODO DEMO LIMITADO")
    with st.sidebar.expander("⭐ ACTIVAR SUSCRIPCIÓN PRO", expanded=True):
        st.markdown("""
        **Planes Disponibles:**
        <div class='plan-box'><b>⏱️ Plan Semanal ($5)</b><br>Acceso completo por 7 días.</div>
        <div class='plan-box'><b>📅 Plan Mensual ($10)</b><br>Acceso completo por 30 días. (Ahorro del 50%)</div>
        <div class='plan-vip'><b>👑 Plan Consultor VIP ($49/mes)</b><br>Acceso + Asesoría Privada y lectura de datos con Darwin López.</div>
        """, unsafe_allow_html=True)
        if os.path.exists("UNA.png"): st.image("UNA.png", caption="Escanea para pagar con Deuna")
        st.markdown("""
        **¿Cómo activar?**
        Envía tu comprobante indicando el plan elegido a:
        <a href="mailto:darwinleonardolopeznarvaez@gmail.com" style="color: #4A90E2; font-weight: bold;">📧 darwinleonardolopeznarvaez@gmail.com</a>
        """, unsafe_allow_html=True)
else:
    st.sidebar.success(f"✨ Licencia de Consultor Activa")

st.sidebar.divider()
if st.sidebar.button("🔄 Iniciar Nuevo Análisis (Limpiar)", on_click=limpiar_pantalla, use_container_width=True):
    st.rerun()

# --- FUNCIÓN ESCUDO ANTI-CAPTURAS (MARCA DE AGUA) ---
def aplicar_sello_seguridad(fig):
    if not ES_PREMIUM:
        fig.add_annotation(
            text="VERSIÓN DEMO<br>OLIMPO ANALYTICS",
            xref="paper", yref="paper",
            x=0.5, y=0.5,
            showarrow=False,
            font=dict(size=35, color="rgba(150, 150, 150, 0.3)"), 
            textangle=-30,
            align="center"
        )
    return fig

# --- ENCABEZADO PRINCIPAL ---
st.markdown("<div class='main-title'>OLIMPO ANALYTICS PRO</div>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #5D6D7E; font-size: 18px;'>Reemplazo Definitivo para SPSS y R - Automatización Científica</p>", unsafe_allow_html=True)
st.divider()

if 'doc_data' not in st.session_state:
    st.session_state.doc_data = {
        'outliers': {'texto': '', 'fig': None},
        'descriptivo': {'texto': '', 'fig': None},
        'bivariado': {'texto': '', 'fig': None},
        'regresion': {'texto': '', 'fig': None},
        'cluster': {'texto': '', 'fig1': None, 'fig2': None}
    }

archivo_subido = st.file_uploader("📂 Carga tu Base de Datos (Formato CSV)", type=["csv"])

if archivo_subido is not None:
    if 'df_trabajo' not in st.session_state:
        df_raw = pd.read_csv(archivo_subido, sep=None, engine='python', dtype=str)
        if not ES_PREMIUM:
            st.session_state.df_trabajo = df_raw.head(min(100, len(df_raw))).copy()
            st.error(f"Modo Demo: Procesando {len(st.session_state.df_trabajo)} registros.")
        else:
            st.session_state.df_trabajo = df_raw.copy()
            
        for col in st.session_state.df_trabajo.columns:
             if st.session_state.df_trabajo[col].str.isnumeric().any() or st.session_state.df_trabajo[col].str.replace('.','',1).str.isnumeric().any():
                  st.session_state.df_trabajo[col] = pd.to_numeric(st.session_state.df_trabajo[col], errors='coerce')

    df_trabajo = st.session_state.df_trabajo

    # --- 🛠️ PASO 1: PREPARACIÓN ---
    st.markdown("<h3 class='section-header'>🛠️ Paso 1: Preparación de la Base de Datos</h3>", unsafe_allow_html=True)
    col_izq, col_der = st.columns(2)
    with col_izq:
        with st.expander("A. Renombrar y Tipo de Variable", expanded=False):
            var_a_renombrar = st.selectbox("1. Selecciona variable:", df_trabajo.columns, key="sel_renombrar")
            nuevo_nombre = st.text_input("2. Nuevo nombre:", value=var_a_renombrar)
            es_num_actual = pd.api.types.is_numeric_dtype(df_trabajo[var_a_renombrar])
            tipo_nuevo = st.radio("3. Tipo:", ["Numérica", "Categórica (Texto)"], index=0 if es_num_actual else 1)
            
            if st.button("Aplicar Cambio", type="primary", use_container_width=True):
                if nuevo_nombre != var_a_renombrar: df_trabajo = df_trabajo.rename(columns={var_a_renombrar: nuevo_nombre})
                if "Numérica" in tipo_nuevo: df_trabajo[nuevo_nombre] = pd.to_numeric(df_trabajo[nuevo_nombre], errors='coerce')
                else: df_trabajo[nuevo_nombre] = df_trabajo[nuevo_nombre].astype(str)
                st.session_state.df_trabajo = df_trabajo
                st.rerun()

    with col_der:
        with st.expander("B. Recodificar (Traducir Códigos)", expanded=False):
            var_recod = st.selectbox("Selecciona variable a traducir:", ["(Elige)"] + list(df_trabajo.columns), key="sel_recod")
            if var_recod != "(Elige)":
                valores_unicos = df_trabajo[var_recod].dropna().unique()
                if len(valores_unicos) <= 25: 
                    diccionario_reemplazo = {}
                    for val in valores_unicos: diccionario_reemplazo[val] = st.text_input(f"'{val}' significa:", value=str(val), key=f"rec_{val}")
                    if st.button("Guardar Traducciones", type="primary", use_container_width=True):
                        df_trabajo[var_recod] = df_trabajo[var_recod].replace(diccionario_reemplazo)
                        st.session_state.df_trabajo = df_trabajo 
                        st.rerun()
                else: st.warning("⚠️ Variable continua. Elige una categórica.")

    st.markdown("#### 👀 Vista Previa de la Base de Datos")
    st.dataframe(df_trabajo.head(5), use_container_width=True)
    st.divider()

    tabs = st.tabs(["🧹 Calidad (Vacíos)", "🕵️ Tratamiento de Outliers", "📊 Univariado", "⚖️ Bivariado", "📈 Regresión", "🤖 Clúster", "📄 Reporte en Word"])
    
    # 🔒 CONFIGURACIÓN DE SEGURIDAD (Oculta botones de descarga nativos de Plotly)
    config_graf = {'displayModeBar': False} if not ES_PREMIUM else {'displayModeBar': True, 'displaylogo': False}

    # --- TAB 1: CALIDAD ---
    with tabs[0]:
        st.markdown("<div class='edu-note'><b>🧠 Metodología Científica:</b> Para rellenar huecos usamos la <b>Mediana</b> porque es una medida <i>robusta</i>. Para texto, usamos la <b>Moda</b>.</div>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            duplicados = df_trabajo.duplicated().sum()
            if duplicados > 0:
                st.error(f"⚠️ **{duplicados}** filas idénticas.")
                if st.button("🗑️ Eliminar Repetidas"):
                    st.session_state.df_trabajo = df_trabajo.drop_duplicates()
                    st.rerun()
            else: st.success("✅ Sin filas repetidas.")
        with c2:
            nulos = (df_trabajo.isna().sum() / len(df_trabajo) * 100).reset_index()
            nulos.columns = ['Variable', '% Faltante']
            st.dataframe(nulos.style.format({'% Faltante': '{:.1f}%'}), use_container_width=True)
        
        if st.toggle("✨ Ejecutar Imputación Científica"):
            for col in df_trabajo.columns:
                if df_trabajo[col].isna().sum() > 0:
                    if pd.api.types.is_numeric_dtype(df_trabajo[col]): df_trabajo[col].fillna(df_trabajo[col].median(), inplace=True)
                    else: df_trabajo[col].fillna(df_trabajo[col].mode()[0], inplace=True)
            st.session_state.df_trabajo = df_trabajo
            st.success("✅ Base purificada.")
            
            # 🔒 BLOQUEO DE DESCARGA DE BASE
            if ES_PREMIUM:
                csv_limpio = df_trabajo.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Descargar Base Depurada (.csv)", data=csv_limpio, file_name='Datos_Limpios_Olimpo.csv', mime='text/csv', type="primary")
            else:
                st.markdown("<div class='premium-lock'>🔒 Descarga de base de datos depurada disponible solo en la versión Premium.</div>", unsafe_allow_html=True)

    # --- TAB 2: OUTLIERS ---
    with tabs[1]:
        st.markdown("<div class='edu-note'><b>🎯 Detección de Outliers (IQR):</b> Caza datos raros que pueden arruinar tus promedios.</div>", unsafe_allow_html=True)
        num_cols = df_trabajo.select_dtypes(include=np.number).columns.tolist()
        if len(num_cols) > 0:
            var_outlier = st.selectbox("Selecciona la variable a inspeccionar:", num_cols)
            Q1 = df_trabajo[var_outlier].quantile(0.25)
            Q3 = df_trabajo[var_outlier].quantile(0.75)
            IQR = Q3 - Q1
            limite_inferior = Q1 - 1.5 * IQR
            limite_superior = Q3 + 1.5 * IQR
            
            outliers_detectados = df_trabajo[(df_trabajo[var_outlier] < limite_inferior) | (df_trabajo[var_outlier] > limite_superior)]
            cantidad_outliers = len(outliers_detectados)
            porcentaje_outliers = (cantidad_outliers / len(df_trabajo)) * 100

            c1, c2 = st.columns([1, 2])
            with c1:
                st.markdown("#### 🚨 Diagnóstico")
                if cantidad_outliers > 0:
                    st.error(f"**{cantidad_outliers}** atípicos detectados.")
                    st.write(f"Representan el **{porcentaje_outliers:.1f}%** de los datos.")
                else: st.success("✅ Variable limpia.")
            
            with c2:
                fig_out = px.box(df_trabajo, x=var_outlier, points="all", title=f"Atípicos en {var_outlier}", color_discrete_sequence=['#E74C3C'])
                fig_out = aplicar_sello_seguridad(fig_out) # Sello de Agua
                st.plotly_chart(fig_out, use_container_width=True, config=config_graf)

            if cantidad_outliers > 0:
                st.markdown("#### 🛠️ Soluciones")
                opcion_outlier = st.radio("Selecciona método:", ["Reemplazar por la Mediana (Conservador)", "Winsorización al P5 y P95 (Técnico)", "Eliminar filas completas (Drástico)"])
                if st.button("Aplicar Solución Seleccionada", type="primary"):
                    if "Mediana" in opcion_outlier:
                        med = df_trabajo[var_outlier].median()
                        df_trabajo.loc[(df_trabajo[var_outlier] < limite_inferior) | (df_trabajo[var_outlier] > limite_superior), var_outlier] = med
                        msg = f"Se reemplazaron {cantidad_outliers} outliers por la mediana ({med:.2f})."
                    elif "Winsorización" in opcion_outlier:
                        p5 = df_trabajo[var_outlier].quantile(0.05)
                        p95 = df_trabajo[var_outlier].quantile(0.95)
                        df_trabajo.loc[df_trabajo[var_outlier] < limite_inferior, var_outlier] = p5
                        df_trabajo.loc[df_trabajo[var_outlier] > limite_superior, var_outlier] = p95
                        msg = f"Valores limitados entre {p5:.2f} y {p95:.2f}."
                    else:
                        df_trabajo = df_trabajo[(df_trabajo[var_outlier] >= limite_inferior) & (df_trabajo[var_outlier] <= limite_superior)]
                        msg = f"Se eliminaron {cantidad_outliers} filas."
                    st.session_state.df_trabajo = df_trabajo
                    st.session_state.doc_data['outliers']['texto'] = f"Tratamiento de Outliers en {var_outlier}: {msg}"
                    st.rerun()
        else: st.warning("⚠️ Sin variables numéricas.")

    # --- TAB 3: UNIVARIADO ---
    with tabs[2]:
        st.markdown("<div class='edu-note'><b>🎯 Univariado:</b> Resumen estadístico de UNA sola variable.</div>", unsafe_allow_html=True)
        v = st.selectbox("Selecciona la variable a diagnosticar:", df_trabajo.columns, key="sel_desc")
        
        if pd.api.types.is_numeric_dtype(df_trabajo[v]):
            stats = df_trabajo[v].describe()
            moda_val = df_trabajo[v].mode()[0] if not df_trabajo[v].mode().empty else "N/A"
            asimetria = df_trabajo[v].skew()
            
            c1, c2, c3, c4, c5 = st.columns(5)
            c1.metric("Promedio (Media)", f"{stats['mean']:.2f}")
            c2.metric("Mediana", f"{stats['50%']:.2f}")
            c3.metric("Moda", f"{moda_val:.2f}" if isinstance(moda_val, (int, float)) else moda_val)
            c4.metric("Mínimo", f"{stats['min']:.2f}")
            c5.metric("Máximo", f"{stats['max']:.2f}")

            with st.expander("Ver Dispersión y Forma (Cuartiles, Asimetría)"):
                col_a, col_b, col_c = st.columns(3)
                col_a.write(f"**Desviación Estándar:** {stats['std']:.2f}")
                col_a.write(f"**Varianza:** {df_trabajo[v].var():.2f}")
                col_b.write(f"**Q1 (25%):** {stats['25%']:.2f}")
                col_b.write(f"**Q3 (75%):** {stats['75%']:.2f}")
                col_c.write(f"**Asimetría:** {asimetria:.2f}")
                col_c.write(f"**Curtosis:** {df_trabajo[v].kurt():.2f}")
            
            texto_sesgo = "sesgo a la derecha" if asimetria > 0 else "sesgo a la izquierda"
            interp_text = f"Promedio de **{v}**: {stats['mean']:.2f}. El 50% de la muestra es menor o igual a {stats['50%']:.2f}. Asimetria ({asimetria:.2f}) indica un {texto_sesgo}."
            st.markdown(f"<div class='interpret-box'><b>🧠 Interpretación:</b> {interp_text}</div>", unsafe_allow_html=True)
            st.session_state.doc_data['descriptivo']['texto'] = interp_text
            
            tipo_graf = st.selectbox("Gráfica Visual:", ["Histograma", "Caja", "Dispersión", "Líneas"], key="graf_desc")
            if tipo_graf == "Histograma": fig_desc = px.histogram(df_trabajo, x=v, marginal="box", color_discrete_sequence=['#5DADE2'])
            elif tipo_graf == "Caja": fig_desc = px.box(df_trabajo, y=v, points="all", color_discrete_sequence=['#F1948A'])
            elif tipo_graf == "Dispersión": fig_desc = px.scatter(df_trabajo, x=df_trabajo.index, y=v, color_discrete_sequence=['#9B59B6'])
            else: fig_desc = px.line(df_trabajo, x=df_trabajo.index, y=v, color_discrete_sequence=['#2ECC71'])
            if tipo_graf not in ["Histograma", "Caja"]: fig_desc.update_layout(xaxis_title="Registro", yaxis_title=v)
            
            fig_desc = aplicar_sello_seguridad(fig_desc) # Sello de Agua
            st.plotly_chart(fig_desc, use_container_width=True, config=config_graf)
            st.session_state.doc_data['descriptivo']['fig'] = fig_desc

        else:
            st.info("📊 Variable Categórica.")
            res = df_trabajo[v].value_counts().reset_index()
            res.columns = [v, 'Conteo']
            tipo_graf_cat = st.radio("Formato:", ["Pastel", "Barras"], horizontal=True, key="graf_cat")
            c1, c2 = st.columns([1, 2])
            with c1: st.dataframe(res)
            with c2: 
                if tipo_graf_cat == "Pastel": fig_desc = px.pie(res, names=v, values='Conteo', hole=0.4, color_discrete_sequence=px.colors.qualitative.Pastel)
                else: fig_desc = px.bar(res, x=v, y='Conteo', color=v, color_discrete_sequence=px.colors.qualitative.Pastel)
                
                fig_desc = aplicar_sello_seguridad(fig_desc) # Sello de Agua
                st.plotly_chart(fig_desc, use_container_width=True, config=config_graf)
                st.session_state.doc_data['descriptivo']['fig'] = fig_desc
                st.session_state.doc_data['descriptivo']['texto'] = f"Frecuencias para la variable categorica {v}."

    # --- TAB 4: BIVARIADO ---
    with tabs[3]:
        st.markdown("<div class='edu-note'><b>🎯 Bivariado:</b> Compara el comportamiento de una variable numérica dividida por los grupos de una variable categórica.</div>", unsafe_allow_html=True)
        num_cols = df_trabajo.select_dtypes(include=np.number).columns.tolist()
        cat_cols = df_trabajo.select_dtypes(exclude=np.number).columns.tolist()
        
        if len(num_cols) > 0 and len(cat_cols) > 0:
            c1, c2 = st.columns(2)
            var_cat = c1.selectbox("Variable Categórica (Agrupador - Eje X):", cat_cols)
            var_num = c2.selectbox("Variable Numérica (A Medir - Eje Y):", num_cols)
            tipo_graf_biv = st.selectbox("Elige la gráfica de comparación:", ["Cajas y Bigotes", "Violín", "Barras de Promedios"])
            
            # --- CORRECCIÓN: SE USAN PALETAS DE COLORES DISCRETAS PARA MANTENER EL COLOR ---
            if "Cajas" in tipo_graf_biv: 
                fig_biv = px.box(df_trabajo, x=var_cat, y=var_num, color=var_cat, title=f"Comparativa de {var_num} según {var_cat}", color_discrete_sequence=px.colors.qualitative.Pastel)
            elif "Violín" in tipo_graf_biv: 
                fig_biv = px.violin(df_trabajo, x=var_cat, y=var_num, color=var_cat, box=True, title=f"Distribución de {var_num} según {var_cat}", color_discrete_sequence=px.colors.qualitative.Pastel)
            else:
                df_agrupado = df_trabajo.groupby(var_cat)[var_num].mean().reset_index()
                fig_biv = px.bar(df_agrupado, x=var_cat, y=var_num, color=var_cat, text_auto='.2f', title=f"Promedio de {var_num} por {var_cat}", color_discrete_sequence=px.colors.qualitative.Pastel)
            
            fig_biv = aplicar_sello_seguridad(fig_biv) # Sello de Agua
            st.plotly_chart(fig_biv, use_container_width=True, config=config_graf)
            interp_biv = f"Analisis Bivariado: Comparacion de la variable '{var_num}' segmentada por los grupos de '{var_cat}'."
            st.markdown(f"<div class='interpret-box'><b>🧠 Nota:</b> {interp_biv}</div>", unsafe_allow_html=True)
            st.session_state.doc_data['bivariado']['texto'] = interp_biv
            st.session_state.doc_data['bivariado']['fig'] = fig_biv
        else: st.warning("⚠️ Necesitas 1 variable Numérica y 1 Categórica para el cruce.")

    # --- TAB 5: REGRESIÓN ---
    with tabs[4]:
        st.markdown("<div class='edu-note'><b>🎯 ¿Para qué sirve?</b> Encuentra la fórmula matemática que conecta dos variables numéricas para predecir el futuro.</div>", unsafe_allow_html=True)
        num_cols = df_trabajo.select_dtypes(include=np.number).columns
        if len(num_cols) >= 2:
            c1, c2 = st.columns(2)
            col_x = c1.selectbox("Variable Independiente (X):", num_cols)
            col_y = c2.selectbox("Variable Dependiente (Y):", num_cols, index=1 if len(num_cols)>1 else 0)
            if col_x != col_y:
                df_mod = df_trabajo[[col_x, col_y]].dropna()
                X = sm.add_constant(df_mod[col_x])
                modelo = sm.OLS(df_mod[col_y], X).fit()
                st.markdown(f"#### 📈 Ecuación: Y = {modelo.params.iloc[0]:.4f} + ({modelo.params.iloc[1]:.4f} * X)")
                st.write(f"**R² (Bondad de ajuste):** {modelo.rsquared:.4f}")
                interp_reg = f"El modelo de Regresion Lineal explica el {modelo.rsquared*100:.1f}% del comportamiento de {col_y} en base a {col_x}. Por cada unidad extra en {col_x}, la proyeccion indica que {col_y} cambiará en {modelo.params.iloc[1]:.4f}."
                st.markdown(f"<div class='interpret-box'><b>🧠 Conclusión:</b> {interp_reg}</div>", unsafe_allow_html=True)
                st.session_state.doc_data['regresion']['texto'] = interp_reg
                
                fig_reg = px.scatter(df_mod, x=col_x, y=col_y, trendline="ols", trendline_color_override="red", title=f"Prediccion: {col_y} vs {col_x}")
                fig_reg = aplicar_sello_seguridad(fig_reg) # Sello de Agua
                st.plotly_chart(fig_reg, use_container_width=True, config=config_graf)
                st.session_state.doc_data['regresion']['fig'] = fig_reg
        else: st.warning("⚠️ Se necesitan 2 variables numéricas.")

    # --- TAB 6: CLUSTERING ---
    with tabs[5]:
        st.markdown("<div class='edu-note'><b>🎯 ¿Qué es el Clustering?</b> Algoritmo IA que agrupa los datos en 'N' grupos basados en similitudes.</div>", unsafe_allow_html=True)
        if not ES_PREMIUM: st.warning("Requiere Licencia Pro para procesar Inteligencia Artificial.")
        else:
            df_num = df_trabajo.select_dtypes(include=np.number).dropna()
            if len(df_num.columns) >= 2:
                n_clus = st.slider("¿Cuántos grupos (clústeres) deseas formar?", 2, 6, 3)
                if st.button("Ejecutar Segmentación IA"):
                    with st.spinner("Entrenando modelo..."):
                        kmeans = KMeans(n_clusters=n_clus, random_state=42, n_init=10)
                        scaler = StandardScaler()
                        datos_escalados = scaler.fit_transform(df_num)
                        df_trabajo.loc[df_num.index, 'Segmento'] = kmeans.fit_predict(datos_escalados)
                        st.session_state.df_trabajo = df_trabajo 
                        st.success(f"✅ Se formaron {n_clus} segmentos.")
                        
                        st.markdown("#### 1. Promedios por Grupo")
                        perfiles = df_trabajo.dropna(subset=['Segmento']).groupby('Segmento')[df_num.columns].mean().reset_index()
                        try: st.dataframe(perfiles.style.background_gradient(cmap='Greens'), use_container_width=True)
                        except: st.dataframe(perfiles, use_container_width=True)
                        
                        st.markdown("#### 2. Visualización Espacial (PCA)")
                        pca = PCA(n_components=2)
                        componentes = pca.fit_transform(datos_escalados)
                        df_pca = pd.DataFrame(data=componentes, columns=['Comp_1', 'Comp_2'], index=df_num.index)
                        df_pca['Segmento'] = df_trabajo.loc[df_num.index, 'Segmento'].astype(str)
                        fig_cluster = px.scatter(df_pca, x='Comp_1', y='Comp_2', color='Segmento', title="Mapa de Nubes de Puntos")
                        st.plotly_chart(fig_cluster, use_container_width=True)
                        st.session_state.doc_data['cluster']['fig1'] = fig_cluster
                        
                        st.markdown("#### 3. Dendrograma Jerárquico")
                        sample_data = datos_escalados[:min(100, len(datos_escalados))]
                        fig_dendro = ff.create_dendrogram(sample_data)
                        fig_dendro.update_layout(width=800, height=500, title="Dendrograma de Similitudes")
                        st.plotly_chart(fig_dendro, use_container_width=True, config=config_graf)
                        st.session_state.doc_data['cluster']['fig2'] = fig_dendro
                        st.session_state.doc_data['cluster']['texto'] = f"El algoritmo segmentó la base en {n_clus} grupos homogeneos, visualizados mediante nubes de puntos y dendrograma."

            else: st.warning("⚠️ Faltan variables numéricas.")

    # --- TAB 7: GENERADOR DE REPORTE EN WORD ---
    with tabs[6]:
        st.markdown("<h3 class='section-header'>📄 Generar Reporte Editable (.docx)</h3>", unsafe_allow_html=True)
        st.markdown("<p class='edu-note'>Exporta todo tu análisis y tus gráficas a color directo a un documento de Word que podrás editar a tu gusto.</p>", unsafe_allow_html=True)
        
        if not ES_PREMIUM: 
            st.markdown("<div class='premium-lock'>🔒 Función exclusiva para Licencias Premium. Por favor, activa un plan para generar y descargar tus reportes.</div>", unsafe_allow_html=True)
        else:
            if st.button("Descargar Reporte en Word", type="primary"):
                with st.spinner("Creando archivo de Word y capturando gráficas..."):
                    try:
                        doc = Document()
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Arial'
                        font.size = Pt(11)
                        
                        titulo = doc.add_heading('REPORTE DE CONSULTORÍA ESTADÍSTICA', level=0)
                        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph('Generado automáticamente por Olimpo Analytics Pro', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph('')
                        
                        doc.add_heading('1. Ficha Técnica y Calidad de Datos', level=1)
                        texto_out = st.session_state.doc_data['outliers']['texto'] if st.session_state.doc_data.get('outliers') else ''
                        doc.add_paragraph(f"Volumen de la muestra analizada: {len(df_trabajo)} registros operativos.\nSe aplicaron protocolos de imputación estadística para tratar datos faltantes.\n{texto_out}")
                        
                        def insertar_en_word(doc_obj, titulo_seccion, data_key, img_keys=['fig']):
                            if st.session_state.doc_data[data_key]['texto'] != '':
                                doc_obj.add_heading(titulo_seccion, level=1)
                                doc_obj.add_paragraph(st.session_state.doc_data[data_key]['texto'])
                                
                                for img_key in img_keys:
                                    figura = st.session_state.doc_data[data_key].get(img_key)
                                    if figura is not None:
                                        try:
                                            temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                                            # Borramos la marca de agua antes de guardar para el PDF del cliente Pro y forzamos el color blanco de fondo
                                            figura.update_layout(template="plotly_white")
                                            figura.layout.annotations = [] 
                                            figura.write_image(temp_img.name, width=800, height=500)
                                            doc_obj.add_picture(temp_img.name, width=Inches(6.0))
                                            doc_obj.add_paragraph('')
                                        except Exception as e:
                                            doc_obj.add_paragraph(f"[No se pudo adjuntar la gráfica. Error: {e}]")

                        insertar_en_word(doc, '2. Análisis Descriptivo Univariado', 'descriptivo')
                        insertar_en_word(doc, '3. Análisis Bivariado (Cruce)', 'bivariado')
                        insertar_en_word(doc, '4. Modelado Predictivo de Regresión', 'regresion')
                        insertar_en_word(doc, '5. Segmentación Algorítmica Multivariada', 'cluster', img_keys=['fig1', 'fig2'])

                        temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        doc.save(temp_doc.name)
                        
                        with open(temp_doc.name, "rb") as file:
                            WordByte = file.read()

                        st.success("✅ ¡El documento de Word está listo!")
                        st.download_button(label="📥 Descargar Reporte en Word (.docx)", data=WordByte, file_name="Reporte_Olimpo.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', type="primary")

                    except Exception as e:
                        st.error(f"❌ Ocurrió un error al generar el Word: {e}")
                        st.info("Asegúrate de instalar las librerías: pip install python-docx kaleido")